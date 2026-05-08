#!/usr/bin/env python3
"""Generate the Agentic RAG Knowledge Service design document as a Word file with embedded diagrams."""

import io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Colour palette ──────────────────────────────────────────────
C_BLUE = "#2563EB"
C_BLUE_LIGHT = "#DBEAFE"
C_GREEN = "#16A34A"
C_GREEN_LIGHT = "#DCFCE7"
C_ORANGE = "#EA580C"
C_ORANGE_LIGHT = "#FFF7ED"
C_PURPLE = "#9333EA"
C_PURPLE_LIGHT = "#F3E8FF"
C_GRAY = "#6B7280"
C_GRAY_LIGHT = "#F3F4F6"
C_RED = "#DC2626"
C_RED_LIGHT = "#FEF2F2"
C_DARK = "#1F2937"
C_WHITE = "#FFFFFF"


def fig_to_bytes(fig, dpi=180):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight", facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


def add_box(ax, x, y, w, h, label, color, text_color="white", fontsize=9, alpha=1.0):
    box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1", facecolor=color, edgecolor="none", alpha=alpha)
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=fontsize,
            color=text_color, fontweight="bold", wrap=True,
            bbox=dict(boxstyle="round,pad=0.05", facecolor="none", edgecolor="none"))


def add_arrow(ax, x1, y1, x2, y2, color=C_GRAY):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="->", color=color, lw=1.5))


# ── Diagram 1: High-Level Architecture ─────────────────────────
def create_architecture_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.set_title("System Architecture Overview", fontsize=14, fontweight="bold", pad=15, color=C_DARK)

    # Consumers row
    consumers = [("AI Agents\n(MCP)", 0.5), ("Dev Tools\n(MCP)", 2.7), ("Chat UI\n(Web)", 4.9), ("REST API\n(HTTP)", 7.1)]
    for label, x in consumers:
        add_box(ax, x, 6.0, 1.8, 0.7, label, C_BLUE, fontsize=8)

    # Auth layer
    add_box(ax, 0.3, 5.0, 8.8, 0.7, "API Gateway / Auth Layer  (JWT + API Key + mTLS + Agent Registry)", C_ORANGE, fontsize=9)

    # Endpoints row
    endpoints = [("MCP\nServer", 0.5, C_PURPLE), ("OpenAI\nCompat", 2.3, C_PURPLE),
                 ("REST\nAPI", 4.1, C_PURPLE), ("Chat\nUI", 5.9, C_PURPLE), ("Admin\nUI", 7.7, C_PURPLE)]
    for label, x, c in endpoints:
        add_box(ax, x, 4.0, 1.4, 0.7, label, c, fontsize=8)

    # Agent orchestrator
    add_box(ax, 0.3, 2.8, 8.8, 0.9, "Agent Orchestrator Core  (LangGraph)\nQuery Classifier  |  Knowledge Registry  |  Retrieval Strategy Router", C_GREEN, fontsize=9)

    # Infrastructure row
    infra = [("Gemma 4 31B\non vLLM\n(A100/H100)", 0.5, C_DARK), ("Qdrant\nVector DB\n+ E5-large", 3.3, C_DARK),
             ("Databases\n(PostgreSQL\netc.)", 6.1, C_DARK)]
    for label, x, c in infra:
        add_box(ax, x, 1.3, 2.4, 1.1, label, c, fontsize=8)

    # Ingestion pipeline
    add_box(ax, 0.3, 0.2, 8.8, 0.8, "Ingestion Pipeline:  Upload -> Parse -> Classify -> Chunk -> Embed -> Store", C_GRAY, fontsize=9)

    # Arrows
    for _, x in consumers:
        add_arrow(ax, x + 0.9, 6.0, x + 0.9, 5.75)
    add_arrow(ax, 4.7, 5.0, 4.7, 4.75)
    add_arrow(ax, 4.7, 4.0, 4.7, 3.75)
    add_arrow(ax, 1.7, 2.8, 1.7, 2.45)
    add_arrow(ax, 4.5, 2.8, 4.5, 2.45)
    add_arrow(ax, 7.3, 2.8, 7.3, 2.45)

    return fig_to_bytes(fig)


# ── Diagram 2: Ingestion Pipeline ──────────────────────────────
def create_ingestion_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Document Ingestion Pipeline", fontsize=14, fontweight="bold", pad=15, color=C_DARK)

    steps = [
        ("Upload\n+ ACL\nTagging", 0.2, 3.5, C_BLUE),
        ("Parser\n(per doc type)\nPDF | Word |\nSpreadsheet |\nTranscript", 2.2, 3.2, C_GREEN),
        ("Classifier\n(Gemma 4)\nAuto-\ncategorize", 4.2, 3.3, C_PURPLE),
        ("Chunker\n512 tokens\nStructure-\naware", 6.2, 3.3, C_ORANGE),
        ("Embed\n+ Store\nE5-large ->\nQdrant", 8.2, 3.3, C_DARK),
    ]

    for label, x, y, color in steps:
        add_box(ax, x, y, 1.6, 1.4, label, color, fontsize=8)

    for i in range(len(steps) - 1):
        x1 = steps[i][1] + 1.6
        x2 = steps[i + 1][1]
        y = steps[i][2] + 0.7
        add_arrow(ax, x1, y, x2, y, C_GRAY)

    # Supported formats below
    formats = [("PDF\n(digital + OCR)", 0.3, 1.5, C_BLUE_LIGHT),
               ("Word\n(.docx)", 2.2, 1.5, C_GREEN_LIGHT),
               ("Spreadsheets\n(.xlsx, .csv)", 4.1, 1.5, C_ORANGE_LIGHT),
               ("Transcripts\n(speaker tagged)", 6.0, 1.5, C_PURPLE_LIGHT),
               ("Databases\n(schema only)", 7.9, 1.5, C_GRAY_LIGHT)]

    for label, x, y, color in formats:
        add_box(ax, x, y, 1.7, 0.8, label, color, text_color=C_DARK, fontsize=7, alpha=0.9)

    ax.text(5.0, 2.6, "Supported Document Formats", ha="center", fontsize=10, fontweight="bold", color=C_GRAY)

    return fig_to_bytes(fig)


# ── Diagram 3: Agentic Query Flow ──────────────────────────────
def create_agent_flow_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(10, 6.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.5)
    ax.axis("off")
    ax.set_title("Agentic Query Flow", fontsize=14, fontweight="bold", pad=15, color=C_DARK)

    steps = [
        ("1. PLAN\nAnalyze query\nIdentify sub-tasks\nClassify query type", 3.5, 5.5, C_BLUE),
        ("2. FILTER\nLoad user AD groups\nBuild ACL filter\nScope to permitted docs", 3.5, 4.3, C_ORANGE),
        ("3. RETRIEVE\n(parallel strategies)", 3.5, 3.1, C_GREEN),
        ("4. EVALUATE\nSufficient context?\nGaps? Contradictions?", 3.5, 1.9, C_PURPLE),
        ("5. SYNTHESIZE\nGenerate answer\nAttach citations", 3.5, 0.7, C_DARK),
    ]

    for label, x, y, color in steps:
        add_box(ax, x, y, 3.0, 0.8, label, color, fontsize=8)

    for i in range(len(steps) - 1):
        add_arrow(ax, 5.0, steps[i][2], 5.0, steps[i + 1][2] + 0.8, C_GRAY)

    # Retrieval strategies branching from step 3
    strategies = [
        ("Vector\nSearch", 0.3, 3.1, C_GREEN),
        ("Text-to-\nSQL", 7.5, 3.5, C_GREEN),
        ("Map-\nReduce", 7.5, 2.7, C_GREEN),
    ]
    for label, x, y, color in strategies:
        add_box(ax, x, y, 1.2, 0.6, label, color, fontsize=7)

    add_arrow(ax, 3.5, 3.4, 1.5, 3.4, C_GREEN)
    add_arrow(ax, 6.5, 3.7, 7.5, 3.7, C_GREEN)
    add_arrow(ax, 6.5, 3.1, 7.5, 3.1, C_GREEN)

    # Re-retrieve loop
    ax.annotate("Gaps?\nRe-retrieve",
                xy=(3.5, 2.6), xytext=(1.0, 2.0),
                fontsize=7, color=C_RED, fontweight="bold",
                arrowprops=dict(arrowstyle="->", color=C_RED, lw=1.2, connectionstyle="arc3,rad=0.3"))

    return fig_to_bytes(fig)


# ── Diagram 4: Auth & Trust Model ──────────────────────────────
def create_auth_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Three-Layer Authentication & Trust Model", fontsize=14, fontweight="bold", pad=15, color=C_DARK)

    # User
    add_box(ax, 0.3, 2.5, 1.8, 1.5, "User\n(Human)\n\nJWT from\nAD/LDAP", C_BLUE, fontsize=8)

    # Calling app/agent
    add_box(ax, 3.3, 2.5, 2.2, 1.5, "Calling App\nor AI Agent\n\n+ mTLS cert\n+ API key\n+ user JWT", C_PURPLE, fontsize=8)

    # RAG API
    add_box(ax, 6.8, 2.5, 2.5, 1.5, "RAG API\n\nValidates:\n1. mTLS (app)\n2. API key (authz)\n3. JWT (user)", C_GREEN, fontsize=8)

    add_arrow(ax, 2.1, 3.25, 3.3, 3.25, C_GRAY)
    add_arrow(ax, 5.5, 3.25, 6.8, 3.25, C_GRAY)

    # Three layers below
    layers = [
        ("Layer 1: mTLS\n\"This is a known,\nregistered application\"", 0.5, 0.5, C_ORANGE),
        ("Layer 2: API Key\n\"This app is authorized\nto use this API\"", 3.5, 0.5, C_RED),
        ("Layer 3: JWT\n\"This request is for\nthis specific user\"", 6.5, 0.5, C_DARK),
    ]
    for label, x, y, color in layers:
        add_box(ax, x, y, 2.8, 1.2, label, color, fontsize=8)

    return fig_to_bytes(fig)


# ── Diagram 5: MCP Knowledge Service ───────────────────────────
def create_mcp_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    ax.set_title("MCP Knowledge Service — Agent-to-Agent Integration", fontsize=14, fontweight="bold", pad=15, color=C_DARK)

    # Consuming agents
    agents = [("HR\nAgent", 0.5, 5.0, C_BLUE), ("Compliance\nAgent", 2.5, 5.0, C_GREEN),
              ("IT Support\nAgent", 4.5, 5.0, C_ORANGE), ("Custom\nAgents", 6.5, 5.0, C_PURPLE)]
    for label, x, y, color in agents:
        add_box(ax, x, y, 1.5, 0.7, label, color, fontsize=8)

    # MCP Server
    add_box(ax, 0.5, 3.3, 8.5, 1.3, "", C_PURPLE_LIGHT, alpha=0.5)
    ax.text(4.75, 4.4, "MCP Server (Knowledge Service)", ha="center", fontsize=11, fontweight="bold", color=C_PURPLE)

    # High-level tools
    add_box(ax, 0.8, 3.5, 3.8, 0.9, "High-Level Tools\n\nask()  |  summarize_topic()\ncompare()", C_PURPLE, fontsize=8)

    # Low-level tools
    add_box(ax, 5.0, 3.5, 3.8, 0.9, "Low-Level Tools\n\nsearch_documents()  |  query_database()\nlookup_document()  |  search_meetings()", C_DARK, fontsize=7)

    # Arrows from agents to MCP
    for _, x, _, _ in agents:
        add_arrow(ax, x + 0.75, 5.0, x + 0.75, 4.65, C_GRAY)

    # Agent orchestrator
    add_box(ax, 1.5, 1.8, 7.0, 0.9, "Agent Orchestrator Core\n(Full agentic RAG pipeline — shared with all endpoints)", C_GREEN, fontsize=9)

    add_arrow(ax, 4.75, 3.3, 4.75, 2.75, C_GRAY)

    # Double-gated access
    add_box(ax, 1.5, 0.4, 3.2, 1.0, "Agent Permissions\n(from agent registry)\nWhich tools & sources\nthe agent can access", C_ORANGE, fontsize=7)
    add_box(ax, 5.3, 0.4, 3.2, 1.0, "User Permissions\n(from JWT / AD groups)\nWhich documents are\nvisible for this request", C_BLUE, fontsize=7)

    ax.text(4.75, 0.15, "Double-Gated Access Control", ha="center", fontsize=9, fontweight="bold", color=C_RED)

    return fig_to_bytes(fig)


# ── Diagram 6: Query Type Strategies ───────────────────────────
def create_query_strategy_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.5)
    ax.axis("off")
    ax.set_title("Query Classification & Retrieval Strategy Routing", fontsize=14, fontweight="bold", pad=15, color=C_DARK)

    # Query classifier
    add_box(ax, 3.5, 4.3, 3.0, 0.8, "Query Classifier\n(Gemma 4)", C_DARK, fontsize=9)

    # Query types
    types = [
        ("Lookup\n\"What does\npolicy 4.2 say?\"", 0.2, 2.5, C_BLUE, "Top-K\nVector Search", 0.2, 1.0, C_BLUE_LIGHT),
        ("Sweep\n\"All questions\nMike asked\"", 2.1, 2.5, C_GREEN, "Map-Reduce\n+ Metadata\nFilter", 2.1, 1.0, C_GREEN_LIGHT),
        ("Analytical\n\"Q3 revenue?\"", 4.0, 2.5, C_ORANGE, "Text-to-SQL\nDirect DB\nQuery", 4.0, 1.0, C_ORANGE_LIGHT),
        ("Cross-ref\n\"Spending vs\npolicy?\"", 5.9, 2.5, C_PURPLE, "Multi-Source\nChained\nRetrieval", 5.9, 1.0, C_PURPLE_LIGHT),
        ("Temporal\n\"What changed\nlast month?\"", 7.8, 2.5, C_RED, "Date Filter\n+ Comparison", 7.8, 1.0, C_RED_LIGHT),
    ]

    for qlabel, qx, qy, qcolor, slabel, sx, sy, scolor in types:
        add_box(ax, qx, qy, 1.7, 1.0, qlabel, qcolor, fontsize=7)
        add_box(ax, sx, sy, 1.7, 0.9, slabel, scolor, text_color=C_DARK, fontsize=7)
        add_arrow(ax, qx + 0.85, qy, qx + 0.85, sy + 0.9, C_GRAY)

    # Arrows from classifier to types
    for _, qx, _, _, _, _, _, _ in types:
        add_arrow(ax, 5.0, 4.3, qx + 0.85, 3.55, C_GRAY)

    return fig_to_bytes(fig)


# ── Diagram 7: Deployment Architecture ─────────────────────────
def create_deployment_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(10, 4.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.5)
    ax.axis("off")
    ax.set_title("Deployment Architecture (On-Premises)", fontsize=14, fontweight="bold", pad=15, color=C_DARK)

    # GPU Node
    add_box(ax, 0.3, 2.5, 2.8, 1.5, "GPU Node\n(A100/H100)\n\nGemma 4 31B\nvia vLLM\n+ E5-large", C_DARK, fontsize=8)

    # App Server
    add_box(ax, 3.6, 2.5, 2.8, 1.5, "App Server\n\nFastAPI\nMCP Server\nIngestion\nPipeline", C_BLUE, fontsize=8)

    # Data Stores
    add_box(ax, 6.9, 2.8, 2.8, 1.2, "Data Stores\n\nQdrant (vectors)\nPostgreSQL (meta)\nElasticsearch (audit)", C_GREEN, fontsize=8)

    # Arrows
    add_arrow(ax, 3.1, 3.25, 3.6, 3.25, C_GRAY)
    add_arrow(ax, 6.4, 3.25, 6.9, 3.25, C_GRAY)

    # Load balancer
    add_box(ax, 1.5, 0.5, 6.5, 1.2, "Supporting Services\n\nLDAP/AD Connector  |  Prometheus + Grafana  |  Admin UI  |  Docker Compose / K8s", C_ORANGE, fontsize=8)

    return fig_to_bytes(fig)


# ── Diagram 8: Implementation Phases ───────────────────────────
def create_phases_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(10, 4))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    ax.axis("off")
    ax.set_title("Implementation Phasing", fontsize=14, fontweight="bold", pad=15, color=C_DARK)

    phases = [
        ("Phase 1\nCore RAG\n\nIngestion + Vector\nSearch + REST API\n+ Chat UI", 0.2, 1.0, C_BLUE),
        ("Phase 2\nAgentic\nOrchestrator\n\nQuery Classification\nMulti-Strategy\nText-to-SQL", 2.6, 1.0, C_GREEN),
        ("Phase 3\nMCP Server\n\nAgent-to-Agent\nIntegration\nmTLS + Registry", 5.0, 1.0, C_PURPLE),
        ("Phase 4\nKnowledge\nLayer\n\nAuto-Categorization\nAdmin UI\nSelf-Organizing", 7.4, 1.0, C_ORANGE),
    ]

    for label, x, y, color in phases:
        add_box(ax, x, y, 2.1, 2.4, label, color, fontsize=8)

    for i in range(len(phases) - 1):
        x1 = phases[i][1] + 2.1
        x2 = phases[i + 1][1]
        add_arrow(ax, x1, 2.2, x2, 2.2, C_DARK)

    ax.text(5.0, 0.5, "Each phase delivers standalone value. Auth present from Phase 1.", ha="center", fontsize=9, fontstyle="italic", color=C_GRAY)

    return fig_to_bytes(fig)


# ── Build the Word Document ─────────────────────────────────────
def build_document():
    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        hs.font.name = "Calibri"

    def add_heading(text, level=1):
        doc.add_heading(text, level=level)

    def add_para(text, bold=False, italic=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p

    def add_bullet(text):
        p = doc.add_paragraph(text, style="List Bullet")
        return p

    def add_image(img_bytes, width=Inches(6.0)):
        doc.add_picture(img_bytes, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Medium Shading 1 Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for r, row_data in enumerate(rows):
            for c, val in enumerate(row_data):
                table.rows[r + 1].cells[c].text = val
        doc.add_paragraph()  # spacing

    # ════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_heading("Agentic RAG\nKnowledge Service", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = add_para("Design & Architecture Document", italic=True)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    date_line = add_para("May 2026")
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    status_line = add_para("Status: Draft")
    status_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS placeholder
    # ════════════════════════════════════════════════════════════
    add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. The Problem",
        "3. Why Agentic RAG (Not Fine-Tuning)",
        "4. System Architecture",
        "5. Document Ingestion Pipeline",
        "6. How the Agent Thinks: Query Flow",
        "7. Security & Access Control",
        "8. MCP Knowledge Service: Agent-to-Agent Integration",
        "9. Technology Stack",
        "10. Deployment Architecture",
        "11. Implementation Phasing",
        "12. Open Questions & Next Steps",
    ]
    for item in toc_items:
        add_para(item)
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════
    add_heading("1. Executive Summary", level=1)

    add_para(
        "This document describes the design of an on-premises Agentic RAG (Retrieval-Augmented Generation) "
        "system that transforms enterprise documents into an intelligent, queryable knowledge service. "
        "The system ingests financial documents, procedure manuals, spreadsheets, meeting transcripts, "
        "and databases, then allows internal employees to ask natural language questions and receive "
        "accurate, cited answers."
    )
    add_para(
        "Beyond serving as a chatbot, the system acts as an MCP (Model Context Protocol) server, "
        "enabling other AI agents across the organization to tap into the knowledge base as a tool. "
        "This positions the system as a foundational knowledge layer for the enterprise AI ecosystem."
    )
    add_para("Key design principles:", bold=True)
    add_bullet("All AI models are US-origin (Gemma 4 31B by Google, E5-large by Microsoft)")
    add_bullet("Runs entirely on-premises on A100/H100 GPUs — no data leaves the network")
    add_bullet("Document-level access control tied to Active Directory groups")
    add_bullet("Agentic retrieval — the system reasons about how to find information, not just keyword matching")
    add_bullet("Serves as an MCP knowledge service for other AI agents, not just a standalone chatbot")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 2. THE PROBLEM
    # ════════════════════════════════════════════════════════════
    add_heading("2. The Problem", level=1)

    add_para(
        "Enterprise knowledge is scattered across hundreds of documents in multiple formats: "
        "financial reports in PDF, procedures in Word, data in spreadsheets, decisions captured in "
        "meeting transcripts, and live data in databases. Employees spend significant time searching "
        "for information, and institutional knowledge is often siloed by department."
    )

    add_heading("What employees need", level=2)
    add_bullet("Quick, accurate answers to questions about company policies and procedures")
    add_bullet("The ability to trace answers back to source documents (citations)")
    add_bullet("Cross-referencing across document types (e.g., 'Does our Q3 spending comply with policy 4.2?')")
    add_bullet("Exhaustive searches across large document sets (e.g., 'What questions did Mike ask in all meetings this month?')")

    add_heading("What the organization needs", level=2)
    add_bullet("Access control — finance docs stay with finance, IT docs stay with IT")
    add_bullet("On-premises deployment — sensitive financial data cannot leave the network")
    add_bullet("An AI knowledge layer that other internal systems and agents can integrate with")
    add_bullet("A system that grows organically as new documents and departments are added")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 3. WHY AGENTIC RAG
    # ════════════════════════════════════════════════════════════
    add_heading("3. Why Agentic RAG (Not Fine-Tuning)", level=1)

    add_para(
        "An initial consideration was LoRA fine-tuning a small model on company documents. "
        "After analysis, this approach was ruled out in favor of agentic RAG for several critical reasons:"
    )

    add_heading("Why not LoRA fine-tuning?", level=2)

    add_table(
        ["Concern", "Fine-Tuning (LoRA)", "Agentic RAG"],
        [
            ["Factual accuracy", "Learns patterns, not facts. Hallucinates specifics.", "Retrieves actual document content. Grounded in source material."],
            ["Citations", "Impossible. Generates from weights, not documents.", "Built-in. Every answer traceable to source."],
            ["Document updates", "Requires retraining on every change.", "Just re-index. No retraining needed."],
            ["Cost of updates", "GPU hours for each retrain cycle.", "Minutes to re-embed changed documents."],
            ["Financial domain risk", "Small models hallucinate financials confidently.", "Only answers from retrieved, verified content."],
        ],
    )

    add_heading("Why agentic RAG specifically?", level=2)

    add_para(
        "Basic RAG follows a simple pattern: embed a query, find similar chunks, generate an answer. "
        "This works for straightforward lookups but fails on complex queries that require:"
    )
    add_bullet("Searching across multiple document types with different retrieval strategies")
    add_bullet("Exhaustive coverage (not just top-K results) for sweep queries")
    add_bullet("Combining database queries with document searches")
    add_bullet("Multi-step reasoning that follows cross-references between documents")

    add_para(
        "An agentic approach adds a reasoning layer that classifies the query type, plans a retrieval "
        "strategy, executes multiple search strategies in parallel, evaluates whether it has enough "
        "context, and can re-retrieve if needed — all before synthesizing an answer."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 4. SYSTEM ARCHITECTURE
    # ════════════════════════════════════════════════════════════
    add_heading("4. System Architecture", level=1)

    add_para(
        "The system is organized in layers: consumers at the top, a shared authentication gateway, "
        "multiple integration endpoints, a core agent orchestrator, and infrastructure services at the bottom. "
        "All endpoints share the same orchestrator core, auth layer, and audit infrastructure."
    )

    add_image(create_architecture_diagram())

    add_heading("Key architectural decisions", level=2)
    add_bullet(
        "Single agent core shared by all endpoints — the MCP server, REST API, OpenAI-compatible API, "
        "and Chat UI all use the same orchestrator. No duplicated logic."
    )
    add_bullet(
        "Four integration points serve different consumers: Chat UI for end users, REST API for custom apps, "
        "OpenAI-compatible API for ecosystem tools, and MCP server for other AI agents."
    )
    add_bullet(
        "The Knowledge Registry is a metadata layer that catalogs all data sources, their types, retrieval "
        "strategies, and access control groups. It helps the agent make informed decisions about where to look."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 5. INGESTION PIPELINE
    # ════════════════════════════════════════════════════════════
    add_heading("5. Document Ingestion Pipeline", level=1)

    add_para(
        "Documents enter the system through an upload UI or a watched folder. Each document passes through "
        "format-specific parsing, automatic classification, structure-aware chunking, and embedding before "
        "being stored in the vector database with rich metadata."
    )

    add_image(create_ingestion_diagram())

    add_heading("How each format is handled", level=2)

    add_table(
        ["Format", "Parser", "Special Handling"],
        [
            ["PDF (digital)", "Unstructured.io", "Direct text extraction, preserves page numbers"],
            ["PDF (scanned)", "Unstructured.io + Tesseract OCR", "OCR converts images to text"],
            ["Word (.docx)", "python-docx via Unstructured.io", "Preserves heading structure for section-aware chunks"],
            ["Spreadsheets", "Unstructured.io", "Stored as text summaries AND queryable tabular data"],
            ["Meeting transcripts", "Custom parser", "Speaker identification, utterance type tagging (question/statement/action item)"],
            ["Databases", "Schema registration only", "Not ingested; queried live via text-to-SQL at runtime"],
        ],
    )

    add_heading("Auto-categorization", level=2)
    add_para(
        "When a new document is uploaded, the LLM (Gemma 4 31B) classifies it against existing categories. "
        "If it matches an existing category, it's assigned automatically. If a new category is detected, "
        "the system proposes the category with a name, description, suggested routing rules, and suggested "
        "ACL groups — then queues it for admin approval before going live."
    )
    add_para(
        "Admin approval is required for new categories because each category carries access control "
        "implications and affects query routing for all users."
    )

    add_heading("Meeting transcript enrichment", level=2)
    add_para(
        "Meeting transcripts receive special processing. Each utterance is tagged with the speaker name "
        "and classified as a question, statement, or action item. This enables powerful metadata-based "
        "queries like 'What questions did Mike ask in all meetings the last 30 days?' — which would be "
        "nearly impossible with naive semantic search alone."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 6. AGENT QUERY FLOW
    # ════════════════════════════════════════════════════════════
    add_heading("6. How the Agent Thinks: Query Flow", level=1)

    add_para(
        "The agent orchestrator is the core reasoning engine. Unlike basic RAG which always does the same "
        "thing (embed query, find top-K chunks, generate), the agentic approach classifies the query type "
        "and selects the best retrieval strategy."
    )

    add_image(create_agent_flow_diagram())

    add_heading("Five-step reasoning process", level=2)
    add_bullet(
        "PLAN — The agent analyzes the query and breaks it into sub-tasks. A complex question like "
        "'What was our Q3 revenue and does it comply with policy 4.2?' becomes two sub-tasks: "
        "one for the database, one for procedure documents."
    )
    add_bullet(
        "FILTER — The user's Active Directory groups are loaded and used to build an ACL filter. "
        "All subsequent retrievals are scoped to only documents the user has permission to see."
    )
    add_bullet(
        "RETRIEVE — The agent executes the appropriate retrieval strategy for each sub-task, "
        "in parallel where possible. Different strategies are used for different query types."
    )
    add_bullet(
        "EVALUATE — The agent reviews the retrieved context. Is it sufficient to answer? "
        "Are there gaps? If so, it reformulates and re-retrieves. Are there contradictions? It flags them."
    )
    add_bullet(
        "SYNTHESIZE — The agent generates a final answer grounded in the retrieved context, "
        "with citations pointing back to specific documents, pages, and sections."
    )

    add_heading("Query type classification", level=2)
    add_para(
        "Not all queries are the same. The agent classifies each query and routes it to the "
        "optimal retrieval strategy:"
    )

    add_image(create_query_strategy_diagram())

    add_heading("The sweep query challenge", level=2)
    add_para(
        "Sweep queries like 'What questions did Mike ask in all meetings the last 30 days?' "
        "are the hardest type for traditional RAG. Standard top-K vector search returns the best "
        "5-10 chunks but misses results scattered across dozens of documents."
    )
    add_para("The agentic approach handles this with a multi-strategy pattern:", bold=True)
    add_bullet("First, narrow the document set using metadata filters (doc_type=transcript, date >= 30 days ago)")
    add_bullet("Run keyword search for speaker-tagged lines ('Mike:' followed by question patterns)")
    add_bullet("Run semantic search with multiple reformulations ('questions raised by Mike', 'Mike asked about', 'Mike wanted to know')")
    add_bullet("For remaining documents, use map-reduce: feed each transcript to the LLM and ask it to extract Mike's questions")
    add_bullet("Merge, deduplicate, and sort results chronologically")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 7. SECURITY & ACCESS CONTROL
    # ════════════════════════════════════════════════════════════
    add_heading("7. Security & Access Control", level=1)

    add_para(
        "With sensitive financial documents and cross-department data, the access control system "
        "must be robust. The design uses a three-layer authentication model that verifies both "
        "the calling application and the end user independently."
    )

    add_image(create_auth_diagram())

    add_heading("Three-layer trust model", level=2)

    add_para("Layer 1: mTLS (Mutual TLS)", bold=True)
    add_para(
        "Each registered application or AI agent receives a client certificate signed by the "
        "internal certificate authority. The RAG API only accepts connections with valid client "
        "certificates. This proves the connection is from a known, registered application."
    )

    add_para("Layer 2: API Key", bold=True)
    add_para(
        "Each application gets a unique API key that enables per-app rate limiting, auditing, "
        "and revocation. This proves the app is authorized to use the API."
    )

    add_para("Layer 3: JWT (JSON Web Token)", bold=True)
    add_para(
        "Users authenticate against LDAP/Active Directory and receive a signed JWT. This token "
        "is passed through with every API call — even when routed through another AI agent. "
        "The RAG API validates the JWT signature independently. It never trusts a bare user_id."
    )

    add_heading("Document-level access control", level=2)
    add_para(
        "Every document is tagged with permitted AD groups at ingestion time. At query time, "
        "the user's AD groups (extracted from their JWT) are used to filter all retrieval operations. "
        "If Mike is in the Finance AD group and Bob is in IT Support, they will get different results "
        "for the same query — Mike sees finance documents, Bob sees IT documents."
    )

    add_heading("Double-gated access for AI agents", level=2)
    add_para(
        "When other AI agents consume the MCP server, access is controlled by both the agent's "
        "permissions AND the user's AD groups. An IT support agent cannot access finance documents "
        "even if the underlying user could — because the agent's permissions restrict which sources "
        "it can query."
    )

    add_heading("Audit trail", level=2)
    add_para("Every query is logged with:")
    add_bullet("Which application or agent made the call")
    add_bullet("Which user the request was on behalf of")
    add_bullet("The full query text")
    add_bullet("Which documents were retrieved and returned")
    add_bullet("Which retrieval strategy was used")
    add_bullet("Timestamp")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 8. MCP KNOWLEDGE SERVICE
    # ════════════════════════════════════════════════════════════
    add_heading("8. MCP Knowledge Service: Agent-to-Agent Integration", level=1)

    add_para(
        "The MCP (Model Context Protocol) server is what transforms this from a standalone chatbot "
        "into a foundational knowledge service. Any AI agent in the organization can use the MCP "
        "server to access enterprise knowledge as a tool — the HR agent, compliance agent, IT support "
        "agent, or any future agent can query the knowledge base."
    )

    add_image(create_mcp_diagram())

    add_heading("Two tiers of tools", level=2)

    add_para("High-level tools (agent-friendly):", bold=True)
    add_para(
        "These run the full agentic RAG pipeline and return complete, cited answers. "
        "A consuming agent just calls ask() with a question and gets back a synthesized result. "
        "The consuming agent doesn't need to understand retrieval strategies — the knowledge "
        "service handles everything."
    )
    add_bullet("ask(question, context, depth) — Full Q&A with citations")
    add_bullet("summarize_topic(topic, time_range, format) — Topic summaries from documents")
    add_bullet("compare(item_a, item_b) — Compare two items across the knowledge base")

    add_para("Low-level tools (precise control):", bold=True)
    add_para(
        "For consuming agents that want to orchestrate their own retrieval logic or need "
        "specific data types."
    )
    add_bullet("search_documents() — Vector similarity search with filters")
    add_bullet("query_database() — Natural language to SQL")
    add_bullet("lookup_document() — Fetch a specific document or section")
    add_bullet("search_meetings() — Speaker and type-filtered transcript search")
    add_bullet("list_sources() — Discover available knowledge sources")

    add_heading("Context passing", level=2)
    add_para(
        "Consuming agents can pass context about their workflow to improve retrieval quality. "
        "For example, an HR onboarding agent asking about PTO policy can include context like "
        "'I am onboarding a new exempt employee in California' — this helps the knowledge service "
        "retrieve the most relevant policy sections."
    )

    add_heading("Async support", level=2)
    add_para(
        "Exhaustive sweep queries may take 30+ seconds. The MCP server supports asynchronous "
        "execution: the client receives a job_id immediately and polls for results. Jobs are "
        "stored server-side with a configurable TTL (default 1 hour)."
    )

    add_heading("Why this matters", level=2)
    add_para(
        "Without the MCP server, every AI system in the organization would need its own copy "
        "of the document corpus, its own ingestion pipeline, and its own retrieval logic. "
        "The MCP knowledge service centralizes this: one source of truth, one set of access controls, "
        "one audit trail — consumed by any agent that needs organizational knowledge."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 9. TECHNOLOGY STACK
    # ════════════════════════════════════════════════════════════
    add_heading("9. Technology Stack", level=1)

    add_para("All AI models are US-origin. All infrastructure is self-hostable on-premises.", bold=True)

    add_table(
        ["Component", "Technology", "Rationale"],
        [
            ["LLM", "Gemma 4 31B (Google)", "US-origin, native tool use, 256K context, strong agentic benchmarks"],
            ["LLM Serving", "vLLM", "Best NVIDIA throughput, continuous batching, OpenAI-compat API"],
            ["Embeddings", "E5-large (Microsoft)", "US-origin, proven retrieval quality, self-hostable"],
            ["Vector DB", "Qdrant", "Self-hosted, Rust-based (fast), metadata filtering for ACLs, MIT license"],
            ["Agent Framework", "LangGraph", "Multi-step agentic flows, tool routing, map-reduce patterns"],
            ["API Layer", "FastAPI (Python)", "Async, auto-generates OpenAPI docs"],
            ["MCP Server", "FastMCP (Python SDK)", "Official MCP SDK, SSE transport"],
            ["Chat UI", "Open WebUI", "Self-hosted, OpenAI-compatible"],
            ["Doc Parsing", "Unstructured.io", "PDF, Word, spreadsheets, OCR; open source, US company"],
            ["Auth", "LDAP + JWT (PyJWT)", "Integrates with existing AD infrastructure"],
            ["Audit Logging", "Elasticsearch or PostgreSQL", "Queryable audit trail for compliance"],
            ["Monitoring", "Prometheus + Grafana", "Standard, self-hosted"],
            ["Orchestration", "Docker Compose / K8s", "Docker Compose to start; Kubernetes for production"],
        ],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 10. DEPLOYMENT
    # ════════════════════════════════════════════════════════════
    add_heading("10. Deployment Architecture", level=1)

    add_image(create_deployment_diagram())

    add_heading("Minimal deployment (start here)", level=2)
    add_bullet("1x A100/H100 node — Gemma 4 31B + E5-large embeddings via vLLM")
    add_bullet("1x application server — Qdrant, FastAPI, ingestion pipeline, MCP server")
    add_bullet("Docker Compose to manage all services")
    add_bullet("Handles approximately 5-10 concurrent users")

    add_heading("Production scale", level=2)
    add_bullet("2+ GPU nodes behind a load balancer for LLM redundancy")
    add_bullet("Dedicated Qdrant cluster (3 nodes) for replication and availability")
    add_bullet("Kubernetes for orchestration and auto-scaling")
    add_bullet("Separate ingestion workers for bulk indexing during off-hours")
    add_bullet("Dedicated Elasticsearch cluster for audit logs")

    add_heading("GPU sizing", level=2)
    add_table(
        ["Workload", "Requirement"],
        [
            ["Gemma 4 31B inference", "1x A100 80GB, ~10-15 tokens/sec per request"],
            ["E5-large embeddings", "Can share GPU or run on CPU"],
            ["Concurrent conversations", "1x A100 handles ~5-10 via continuous batching"],
            ["Bulk ingestion", "Can share inference GPU or use separate GPU during off-hours"],
        ],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 11. IMPLEMENTATION PHASING
    # ════════════════════════════════════════════════════════════
    add_heading("11. Implementation Phasing", level=1)

    add_para(
        "The system is designed to be built incrementally, with each phase delivering standalone value. "
        "This reduces risk and allows the team to learn from real usage before adding complexity."
    )

    add_image(create_phases_diagram())

    add_table(
        ["Phase", "Scope", "Delivers"],
        [
            ["Phase 1", "Core RAG: ingestion + vector search + REST API + Chat UI",
             "Users can upload docs and ask questions with citations"],
            ["Phase 2", "Agentic orchestrator: query classification, multi-strategy retrieval, text-to-SQL",
             "Smarter answers, database integration, sweep queries"],
            ["Phase 3", "MCP server + agent-to-agent integration",
             "Other AI systems can use the knowledge base as a tool"],
            ["Phase 4", "Knowledge layer auto-categorization + admin UI",
             "Self-organizing corpus, category management"],
        ],
    )

    add_para(
        "Authentication (JWT + API key) is present from Phase 1. mTLS and the agent registry "
        "are added in Phase 3 when external agents connect.",
        italic=True,
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 12. OPEN QUESTIONS
    # ════════════════════════════════════════════════════════════
    add_heading("12. Open Questions & Next Steps", level=1)

    add_para("The following items need stakeholder input before implementation begins:")

    questions = [
        ("AD group structure", "What existing AD groups map to document access? Needs input from IT/security team."),
        ("Database schemas", "Which databases should be connected? What are their access patterns?"),
        ("Meeting transcript format", "What tool generates transcripts? Do they already have speaker tags?"),
        ("Admin UI scope", "Build custom or use an existing internal tools framework?"),
        ("Backup & disaster recovery", "Strategy for vector DB and metadata store."),
        ("Model update cadence", "Process for evaluating and upgrading to newer model versions."),
    ]

    for title, desc in questions:
        p = doc.add_paragraph()
        run_title = p.add_run(f"{title}: ")
        run_title.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    add_heading("Next steps", level=2)
    add_bullet("Review and approve this design document")
    add_bullet("Resolve open questions with stakeholders")
    add_bullet("Create detailed implementation plan for Phase 1")
    add_bullet("Set up development environment with GPU access")
    add_bullet("Begin Phase 1 implementation")

    # ── Save ──
    output_path = "/Users/michaelmulkey/Documents/Repositories/rag/Agentic_RAG_Knowledge_Service_Design.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Document saved to: {path}")
