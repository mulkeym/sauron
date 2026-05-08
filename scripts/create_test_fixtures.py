import os
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF

os.makedirs("test_fixtures", exist_ok=True)

pdf = FPDF()
pdf.add_page()
pdf.set_font("Helvetica", size=12)
pdf.cell(200, 10, text="Finance Policy Document", new_x="LMARGIN", new_y="NEXT")
pdf.cell(200, 10, text="Section 4.2: Expense Reporting", new_x="LMARGIN", new_y="NEXT")
pdf.cell(200, 10, text="All expenses over $500 require manager approval.", new_x="LMARGIN", new_y="NEXT")
pdf.cell(200, 10, text="Receipts must be submitted within 30 days.", new_x="LMARGIN", new_y="NEXT")
pdf.output("test_fixtures/sample.pdf")

doc = Document()
doc.add_heading("IT Runbook: Server Restart Procedure", level=1)
doc.add_heading("Step 1: Pre-checks", level=2)
doc.add_paragraph("Verify no active deployments are running.")
doc.add_paragraph("Check the monitoring dashboard for anomalies.")
doc.add_heading("Step 2: Restart", level=2)
doc.add_paragraph("SSH into the server and run: sudo systemctl restart app-server")
doc.save("test_fixtures/sample.docx")

wb = Workbook()
ws = wb.active
ws.title = "Q3 Budget"
ws.append(["Department", "Budget", "Spent", "Remaining"])
ws.append(["Engineering", 500000, 420000, 80000])
ws.append(["Marketing", 300000, 290000, 10000])
ws.append(["Finance", 200000, 150000, 50000])
wb.save("test_fixtures/sample.xlsx")

with open("test_fixtures/sample_transcript.txt", "w") as f:
    f.write("Meeting: Engineering Standup\n")
    f.write("Date: 2026-04-10\n")
    f.write("---\n")
    f.write("Mike: Are we on track for the Q2 release?\n")
    f.write("Sarah: Yes, but the API migration is behind schedule.\n")
    f.write("Mike: What's blocking the API migration?\n")
    f.write("Sarah: We're waiting on the new auth library to be approved.\n")
    f.write("Bob: I can help with testing once it's ready.\n")

print("Test fixtures created.")
