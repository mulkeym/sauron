import pytest
from pathlib import Path
from src.ingestion.parser import parse_document, ParsedDocument

FIXTURES = Path(__file__).parent.parent.parent / "test_fixtures"


def test_parse_pdf():
    result = parse_document(FIXTURES / "sample.pdf")
    assert isinstance(result, ParsedDocument)
    assert result.doc_type == "pdf"
    assert "expense" in result.text.lower() or "finance" in result.text.lower()
    assert result.filename == "sample.pdf"


def test_parse_docx():
    result = parse_document(FIXTURES / "sample.docx")
    assert result.doc_type == "docx"
    assert "server restart" in result.text.lower() or "runbook" in result.text.lower()


def test_parse_docx_captures_table_story_text(tmp_path):
    """Illustrated kids books often put body copy in tables, not body paragraphs."""
    from docx import Document

    path = tmp_path / "story.docx"
    doc = Document()
    doc.add_paragraph("Little Axel")  # title in a normal paragraph
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = (
        "It was a quiet morning in the forest. The little fox sat under a maple tree."
    )
    doc.save(path)

    result = parse_document(path)
    assert result.doc_type == "docx"
    assert "quiet morning in the forest" in result.text.lower()
    assert "little fox" in result.text.lower()


def test_parse_xlsx():
    result = parse_document(FIXTURES / "sample.xlsx")
    assert result.doc_type == "xlsx"
    assert "engineering" in result.text.lower() or "budget" in result.text.lower()


def test_parse_transcript():
    result = parse_document(FIXTURES / "sample_transcript.txt")
    assert result.doc_type == "transcript"
    assert len(result.utterances) > 0
    mike_questions = [u for u in result.utterances if u.speaker == "Mike" and u.utterance_type == "question"]
    assert len(mike_questions) == 2


def test_parse_unsupported_format():
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
        f.write(b"some content")
        path = Path(f.name)
    with pytest.raises(ValueError, match="Unsupported"):
        parse_document(path)
    path.unlink()
