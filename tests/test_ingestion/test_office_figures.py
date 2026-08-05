"""Word/Excel embedded image region collection and text enrich."""
from __future__ import annotations

import io
from pathlib import Path
from unittest.mock import patch

import pytest
from PIL import Image


def _png_bytes(w=400, h=300, color=(40, 80, 160)) -> bytes:
    img = Image.new("RGB", (w, h), color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _make_docx_with_image(path: Path, png: bytes) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("Network Overview", level=1)
    doc.add_paragraph("This document describes the SD-WAN portal architecture.")
    # write temp image file for python-docx
    img_path = path.parent / "_tmp_fig.png"
    img_path.write_bytes(png)
    try:
        doc.add_picture(str(img_path), width=Inches(4))
    finally:
        img_path.unlink(missing_ok=True)
    doc.add_paragraph("See figure above for components.")
    doc.save(str(path))


def _make_xlsx_with_image(path: Path, png: bytes) -> None:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "Budget"
    ws["A2"] = 100
    img_path = path.parent / "_tmp_xl.png"
    img_path.write_bytes(png)
    try:
        xl_img = XLImage(str(img_path))
        ws.add_image(xl_img, "C2")
        wb.save(str(path))
    finally:
        img_path.unlink(missing_ok=True)


def test_extract_image_regions_docx(tmp_path):
    from src.ingestion.figure_extract import extract_image_regions_docx

    docx_path = tmp_path / "with_fig.docx"
    _make_docx_with_image(docx_path, _png_bytes(500, 400))
    regions = extract_image_regions_docx(docx_path)
    assert len(regions) >= 1
    assert regions[0].source == "docx"
    assert regions[0].width >= 400
    assert regions[0].image_bytes[:8] == b"\x89PNG\r\n\x1a\n"


def test_docx_parser_preserves_figure_anchor_and_context(tmp_path):
    from docx import Document
    from docx.shared import Inches
    from src.ingestion.parser import parse_document

    docx_path = tmp_path / "anchored.docx"
    img_path = tmp_path / "anchor.png"
    img_path.write_bytes(_png_bytes(500, 400))
    doc = Document()
    doc.add_heading("Network Architecture", level=1)
    doc.add_paragraph("Traffic enters through the branch gateway.")
    doc.add_picture(str(img_path), width=Inches(4))
    doc.inline_shapes[0]._inline.docPr.set("descr", "Redundant controller topology")
    doc.add_paragraph("Figure 4. Controller topology", style="Caption")
    doc.add_paragraph("Controllers are deployed in separate zones.")
    doc.save(str(docx_path))

    parsed = parse_document(docx_path)
    kinds = [block.block_type for block in parsed.blocks]
    figure_at = kinds.index("figure")
    assert kinds[figure_at - 1] == "paragraph"
    assert kinds[figure_at + 1] == "paragraph"
    placement = parsed.blocks[figure_at].figure
    assert placement is not None
    assert placement.figure_id == "fig-001"
    assert placement.section_path == ["Network Architecture"]
    assert placement.caption == "Figure 4. Controller topology"
    assert placement.alt_text == "Redundant controller topology"
    assert placement.previous_text == "Traffic enters through the branch gateway."
    assert placement.following_text == "Controllers are deployed in separate zones."


def test_repeated_docx_image_keeps_both_placements(tmp_path):
    from docx import Document
    from docx.shared import Inches
    from src.ingestion.figure_extract import extract_image_regions_docx

    docx_path = tmp_path / "repeated.docx"
    img_path = tmp_path / "repeat.png"
    img_path.write_bytes(_png_bytes(500, 400))
    doc = Document()
    doc.add_heading("Primary", level=1)
    doc.add_picture(str(img_path), width=Inches(3))
    doc.add_heading("Recovery", level=1)
    doc.add_picture(str(img_path), width=Inches(3))
    doc.save(str(docx_path))

    regions = extract_image_regions_docx(docx_path)
    assert [r.figure_id for r in regions] == ["fig-001", "fig-002"]
    assert regions[0].content_hash == regions[1].content_hash
    assert regions[0].section_path == ["Primary"]
    assert regions[1].section_path == ["Recovery"]


def test_extract_image_regions_xlsx(tmp_path):
    from src.ingestion.figure_extract import extract_image_regions_xlsx

    xlsx_path = tmp_path / "with_fig.xlsx"
    _make_xlsx_with_image(xlsx_path, _png_bytes(480, 320))
    regions = extract_image_regions_xlsx(xlsx_path)
    assert len(regions) >= 1
    assert regions[0].source == "xlsx"
    assert regions[0].width >= 300


def test_enrich_text_with_figures_docx(tmp_path, monkeypatch):
    from src.config import settings
    from src.ingestion.figure_extract import (
        enrich_text_with_figures, FigureRecord, ProseBlock,
    )

    monkeypatch.setattr(settings, "figure_extraction_enabled", True)
    docx_path = tmp_path / "net.docx"
    _make_docx_with_image(docx_path, _png_bytes(600, 400))

    fake_prose = ProseBlock(
        text="[Figure p.1 — network]\nIdentifiers:\n- portal.example.com\n[/Figure]",
        page=0,
    )

    def fake_process(regions, **kwargs):
        from src.ingestion.figure_extract import FigureEnrichmentResult
        region = regions[0]
        return FigureEnrichmentResult(
            prose_blocks=[fake_prose],
            figure_records=[FigureRecord(
                figure_id=region.figure_id,
                description=fake_prose.text,
                kind="network",
                body_index=region.body_index,
                section_path=region.section_path,
                previous_text=region.previous_text,
                following_text=region.following_text,
            )],
            table_grids=[],
            figures_seen=len(regions),
            figures_used=1,
        )

    with patch("src.ingestion.figure_extract.process_image_regions", side_effect=fake_process):
        text, grids = enrich_text_with_figures(
            docx_path, "Body paragraph about the network.",
        )
    assert "SD-WAN portal architecture" in text
    assert "## Embedded figures" not in text
    assert "portal.example.com" in text
    assert text.index("SD-WAN portal architecture") < text.index("portal.example.com")
    assert text.index("portal.example.com") < text.index("See figure above")
    assert grids == []


def test_enrich_text_disabled(tmp_path, monkeypatch):
    from src.config import settings
    from src.ingestion.figure_extract import enrich_text_with_figures

    monkeypatch.setattr(settings, "figure_extraction_enabled", False)
    docx_path = tmp_path / "x.docx"
    _make_docx_with_image(docx_path, _png_bytes())
    text, grids = enrich_text_with_figures(docx_path, "hello")
    assert text == "hello"
    assert grids == []


def test_dispatch_for_path(tmp_path):
    from src.ingestion.figure_extract import extract_image_regions_for_path

    docx_path = tmp_path / "a.docx"
    _make_docx_with_image(docx_path, _png_bytes(400, 300))
    regs = extract_image_regions_for_path(docx_path)
    assert len(regs) >= 1

    xlsx_path = tmp_path / "b.xlsx"
    _make_xlsx_with_image(xlsx_path, _png_bytes(400, 300))
    regs2 = extract_image_regions_for_path(xlsx_path)
    assert len(regs2) >= 1
